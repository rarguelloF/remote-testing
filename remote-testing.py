#!/usr/bin/python

import sys

import optparse
import getpass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

from paramiko import SSHClient, SSHException, AuthenticationException, AutoAddPolicy
import socket
import select

from scapy.all import send, IP, ICMP

COLOURS = {
    'red': (255, 0, 0),
    'green': (0, 255, 0),
    'blue': (0, 0, 255)
}

class sshConnection():

    def __init__(self, host, user, password):
        self.host = host
        self.user = user
        self.password = password
        self.connection = None

    def connect(self):
        try:
            print "[+] Estableciendo conexion SSH con %s@%s...." % (self.user, self.host)
            self.connection = SSHClient()
            self.connection.set_missing_host_key_policy(AutoAddPolicy())
            self.connection.connect(
                self.host, 
                username = self.user, 
                password = self.password, 
                allow_agent = False, 
                look_for_keys = False
            )
        except AuthenticationException, e:
            print "[+] Error de autenticacion: " ,e
            return 1
        except SSHException, e:
            print "[+] Error en SSH: " , e
            return 2
        except socket.error, e:
            print "[+] Error de conexion: ", e
            return 3
        else: 
            print "[+] Conexion establecida."
            return 0

    def send_command(self, command):
        if not self.connection:
            print '[+] No hay ninguna conexion abierta.'
            return 1
        else:
            cmd_output = ''
            # Send the command (non-blocking)
            stdin, stdout, stderr = self.connection.exec_command(command)

            # Wait for the command to terminate
            while not stdout.channel.exit_status_ready():
                # Only print data if there is data to read in the channel
                if stdout.channel.recv_ready():
                    rl, wl, xl = select.select([stdout.channel], [], [], 0.0)
                    if len(rl) > 0:
                        # Print data from stdout
                        cmd_output+=stdout.channel.recv(1024)
            return cmd_output

    def disconnect(self):
        if self.connection:
            self.connection.close()
            return 0
        else:
            print '[+] No hay ninguna conexion abierta.'
            return 1

###########################################
#                TESTS                    #
###########################################

def custom_test(conn, doc):
    '''
    Prueba generica para introducir cualquier otro comando.
    '''
    #Se envia para que muestre la salida de una vez y no haya que pulsar espacio
    resp = conn.send_command('terminal length 0')
    print resp
    cmd = raw_input("Introduzca comando: \n")
    resp = conn.send_command(cmd)
    print resp
    return resp
     
def OSPF_test(conn, doc):
    write_to_doc(doc, "Comprobar OSPF", 'tituloprueba', 'garamond', 16, 
        {'all_caps': True})
    write_to_doc(doc, "OBJETIVO", 'tituloprueba2', 'garamond', 12, 
        {'all_caps': True, 'bold': True, 'underline': True})
    write_to_doc(doc, "Comprobar que la sesion de OSPF se ha establecido", 'normal2', 'garamond', 12)
    write_to_doc(doc, "Pasos a realizar", 'tituloprueba2', 'garamond', 12,
        {'all_caps': True, 'bold': True, 'underline': True})
    write_to_doc(doc, "1 Introducir el comando show ip ospf neighbor y comprobar que el estado es FULL", 'normal2', 'garamond', 12)

    resp = conn.send_command('show ip ospf neighbor')
    write_to_doc(doc, resp, 'normal1', 'courier new', 9)
    print resp

    write_to_doc(doc, "Conclusiones", 'tituloprueba2', 'garamond', 12, 
        {'bold': True, 'underline': True, 'all_caps': True})
    if 'FULL' in resp:
        write_to_doc(doc, "La sesion se ha establecido correctamente", 'normal2', 'garamond', 12)
        write_to_doc(doc, "Resultado", 'tituloprueba2', 'garamond', 12, 
            {'bold': True, 'underline': True, 'all_caps': True})
        write_to_doc(doc, "OK", 'ok', 'garamond', 12, 
            {'bold': True, 'color': 'green'})
    else:
        write_to_doc(doc, "La sesion no se ha establecido correctamente", 'normal2', 'garamond', 12)
        write_to_doc(doc, "Resultado", 'tituloprueba2', 'garamond', 12, 
            {'bold': True, 'underline': True, 'all_caps': True})
        write_to_doc(doc, "NOK", 'ok', 'garamond', 12, 
            {'bold': True, 'color': 'red'})

    return 0
     
def Comprobar_Contador_ACL(conn, doc):
    write_to_doc(doc, "Comprobar_Contador_ACL", 'tituloprueba', 'garamond', 16, 
        {'all_caps': True})
    write_to_doc(doc, "OBJETIVO", 'tituloprueba2', 'garamond', 12, 
        {'all_caps': True, 'bold': True, 'underline': True})
    write_to_doc(doc, "Comprobar que la ACL matchea el trafico", 'normal2', 'garamond', 12)
    write_to_doc(doc, "Pasos a realizar", 'tituloprueba2', 'garamond', 12,
        {'all_caps': True, 'bold': True, 'underline': True})
    write_to_doc(doc, "Enviar trafico con scapy", 'normal2', 'garamond', 12)
    
    src = raw_input("Introduzca IP origen \n")
    dst = raw_input("Introduzca IP destino \n")

    write_to_doc(doc, "packet = IP(src=%s, dst=%s)/ICMP()" % (src, dst), 'normal2', 'garamond', 12)
    write_to_doc(doc, "send(packet*3)", 'normal2', 'garamond', 12)

    packet = IP(src=src, dst=dst)/ICMP()
    send(packet*3)

    #Una vez enviado el trafico enviamos el comando show ip access-list a la consola del router
    resp = conn.send_command('show ip ospf neighbor')
    write_to_doc(doc, resp, 'normal1', 'courier new', 9)
    print resp

    write_to_doc(doc, "Conclusiones", 'conclusiones', 'garamond', 12, 
        {'bold': True, 'underline': True, 'all_caps': True})
    if '3' in resp:
        write_to_doc(doc, "La ACL matchea correctamente los paquetes enviados", 'normal2', 'garamond', 12)
        write_to_doc(doc, "Resultado", 'conclusiones', 'garamond', 12, 
            {'bold': True, 'underline': True, 'all_caps': True})
        write_to_doc(doc, "OK", 'ok', 'garamond', 12, 
            {'bold': True, 'color': 'green'})
    else:
        #Si no encuentra la palabra FULL
        write_to_doc(doc, "La ACL no matchea correctamente los paquetes enviados", 'normal2', 'garamond', 12)
        write_to_doc(doc, "Resultado", 'conclusiones', 'garamond', 12, 
            {'bold': True, 'underline': True, 'all_caps': True})
        write_to_doc(doc, "NOK", 'ok', 'garamond', 12, 
            {'bold': True, 'color': 'red'})

    return 0

###########################################
#           OTRAS FUNCIONES               #
###########################################


def read_config_file(filename):
    '''
    Parsea el fichero de configuracion especificado con la
    siguiente sintaxis:

    user1@host1 password1
    user2@host2 password2
    ...
    '''
    hosts = []
    with open(filename, 'r') as f:
        text = f.read()
        for line in text.split('\n'):
            try:
                host, pwd = line.split(' ')
                user, ip = host.split('@')
            except:
                if line != '':
                    print 'Bad line: %s' % line
            else:
                hosts.append((ip, user, pwd))
    return hosts

def write_to_doc(document, text, style_name, font_name, font_size, opts=None):
    doc_styles = document.styles
    try:
        #intentamos encontrar el estilo en el documento para ver si antes existia
        doc_charstyle = doc_styles[style_name]
    except:
        #si no existia, lo anadimos
        doc_charstyle = doc_styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        doc_font = doc_charstyle.font
        doc_font.size = Pt(font_size)
        doc_font.name = font_name
        if opts:
            for opt, value in opts.iteritems():
                if opt == 'color':
                    r,g,b = COLOURS[value]
                    doc_font.color.rgb = RGBColor(r,g,b)
                if opt == 'bold' and value == True:
                    doc_font.bold = True
                if opt == 'underline' and value == True:
                    doc_font.underline = True
                if opt == 'all_caps' and value == True:
                    doc_font.all_caps = True

    lines = text.split('\n')
    for line in lines:
        document.add_paragraph(line, style = style_name)

    return document

def select_test(conn, doc):
    dict_tests = {
        '1': custom_test,
        '2': OSPF_test,
        '3': Comprobar_Contador_ACL,
    }
    print "PRUEBAS"
    print "1) Elegir comando"
    print "2) Comprobar sesion OSPF"
    print "3) Comprobar contador ACL"

    option = raw_input('Selecciona el numero de prueba a realizar: ')

    try:
        dict_tests[option](conn, doc)
    except KeyError as e:
        print 'Opcion no existe.'

def parse_program_options():
    parser = optparse.OptionParser(usage='Usage: remote-tests.py [options]')

    parser.add_option('-f', help='''Si se especifica, recoge la informacion de los hosts 
        a testear del fichero dado.''', dest='hosts_file', default=None, action='store')
    parser.add_option('-o', help='''Nombre del .doc de salida.''', 
        dest='doc_name', default='Prueba', action='store')

    (opts, args) = parser.parse_args()

    hosts = None
    if opts.hosts_file:
        try:
            hosts = read_config_file(opts.hosts_file)
        except Exception as e:
            print 'Error abriendo el fichero de hosts: ', e

    return (hosts, opts.doc_name)


def main():
    hosts, doc_name = parse_program_options()

    doc = Document()
    doc.add_heading('PRUEBA', 0)

    if hosts:
        for host in hosts:
            ip, user, pwd = host
            ssh_conn = sshConnection(ip, user, pwd)
            code = ssh_conn.connect()
            if code == 1:
                i = 1
                while code == 1 and i<=2:
                    pwd = getpass.getpass("Introduzca password para %s@%s\n" % (user,ip))
                    code = conn1.connect()
                    i+=1
            if code == 0:
                doc.add_heading('Pruebas sobre %s' % ip, level=1)
                try:
                    while(1):
                        select_test(ssh_conn, doc)
                except KeyboardInterrupt:
                    break

            ssh_conn.disconnect()
    else:
        while(1):
            print 'Press CTRL+C to stop.'
            ip = raw_input("Introduzca IP remota \n")
            user = raw_input("Introduzca el usuario \n")
            pwd = getpass.getpass("Introduzca password \n")

            ssh_conn = sshConnection(ip, user, pwd)
            code = ssh_conn.connect()
            if code == 1:
                i = 1
                while code == 1 and i<=2:
                    pwd = getpass.getpass("Introduzca password para %s@%s\n" % (user,ip))
                    code = ssh_conn.connect()
                    i+=1
            if code == 0:
                doc.add_heading('Pruebas sobre %s' % ip, level=1)
                try:
                    while(1):
                        select_test(ssh_conn, doc)
                except KeyboardInterrupt:
                    break

            ssh_conn.disconnect()

    doc.add_page_break()
    doc.save('%s.docx' % doc_name)

if __name__ == '__main__':
    main()
