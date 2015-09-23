#!/usr/bin/python

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

COLOURS = {
    'red': (255, 0, 0),
    'green': (0, 255, 0),
    'blue': (0, 0, 255)
}

DOC_NAME = 'DEFAULT_NAME'

def create(title):
    global DOC_NAME = title

    document = Document()
    document.add_heading(DOC_NAME, 0)
    document.save('%s.docx' % DOC_NAME)
    return document

def end_doc(document):
    document.add_page_break()
    document.save('%s.docx' % DOC_NAME)

def write(document, text, style_name, font_name, font_size, opts=None):
    doc_styles = document.styles
    try:
        #intentamos encontrar el estilo en el documento para ver si antes existia
        doc_charstyle = doc_styles[style_name]
    except:
        #si no existia, lo anadimos
        doc_charstyle = doc_styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        doc_font = doc_charstyle.font
        doc_font.size = Pt(font_size)
        doc_font.name = font_name
        if opts:
            for opt, value in opts.iteritems():
                if opt == 'color':
                    r,g,b = COLOURS[value]
                    doc_font.color.rgb = RGBColor(r,g,b)
                if opt == 'bold' and value == True:
                    doc_font.bold = True
                if opt == 'underline' and value == True:
                    doc_font.underline = True
                if opt == 'all_caps' and value == True:
                    doc_font.all_caps = True

    lines = text.split('\n')
    for line in lines:
        document.add_paragraph(line, style = style_name)

    document.save('%s.docx' % DOC_NAME)
    return document